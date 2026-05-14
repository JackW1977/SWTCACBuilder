"""
parser.py – Document parser for SW Test Case Acceptance Criteria Builder.

Responsibilities:
- Open a .docx file
- Find the requirements table (columns: Requirement # / Description / Acceptance Criteria)
- Find the specifications table (columns: Specification # / Description / Acceptance Criteria)
- Find the Section 9 test-procedure table (columns: Step ID / Procedure Step Description /
  Expected Result of Successful Procedure Step / Notes)
- Extract requirement numbers from the Notes column of each procedure step
- Return structured ParseResult

Matching logic
--------------
The Notes column uses short-form IDs like "SRS-95416", "UIS-112369", or "SDS-112455".
The specification table uses long-form IDs like "TBR-UIS-112369" or "TBR-SDS-112455".
`numbers_match()` normalises both sides and checks whether one string is a suffix of the
other, so all combinations resolve correctly.
"""

import re
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.table import _Cell

# Regex that matches SRS-XXXXX, TBR-UIS-XXXXX, TBR-SDS-XXXXX, UIS-XXXXX, SDS-XXXXX
REQ_NUMBER_RE = re.compile(
    r'\b(?:TBR-)?(?:SRS|UIS|SDS)-\d{4,6}\b',
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class RequirementRow:
    number: str          # e.g. "SRS-95416" or "TBR-UIS-112369"
    description: str
    existing_ac: str     # whatever is already in the Acceptance Criteria cell


@dataclass
class ProcedureStep:
    step_id: str
    description: str
    expected_result: str
    notes: str                         # raw text of Notes cell
    referenced_numbers: List[str] = field(default_factory=list)


@dataclass
class ParseResult:
    requirements: List[RequirementRow] = field(default_factory=list)
    specifications: List[RequirementRow] = field(default_factory=list)
    procedure_steps: List[ProcedureStep] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _cell_text(cell: _Cell, sep: str = "\n") -> str:
    """Return all paragraph text in a cell, joined by *sep*."""
    return sep.join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _table_header_row(table) -> List[str]:
    """Return the lowercased text of every cell in the first row."""
    if not table.rows:
        return []
    return [_cell_text(c, " ").lower() for c in table.rows[0].cells]


def _find_col(headers: List[str], *keywords) -> int:
    """Return index of first column whose header contains ALL keywords (case-insensitive)."""
    for i, h in enumerate(headers):
        if all(kw.lower() in h for kw in keywords):
            return i
    return -1


def _find_col_any(headers: List[str], *keywords) -> int:
    """Return index of first column whose header contains ANY of the keywords."""
    for i, h in enumerate(headers):
        if any(kw.lower() in h for kw in keywords):
            return i
    return -1


def _extract_numbers(text: str) -> List[str]:
    """Find all requirement/spec numbers in a block of text."""
    return REQ_NUMBER_RE.findall(text)


def normalize(num: str) -> str:
    """Uppercase, strip whitespace/invisible chars for comparison."""
    return re.sub(r'[\s​ ]+', '', num.strip()).upper()


def numbers_match(a: str, b: str) -> bool:
    """
    Return True if two requirement/spec IDs refer to the same item.

    Handles prefix variants:
      "UIS-112369" matches "TBR-UIS-112369"
      "SDS-112455" matches "TBR-SDS-112455"
    """
    na, nb = normalize(a), normalize(b)
    if na == nb:
        return True
    # One is a suffix of the other (handles TBR- prefix)
    if na.endswith(nb) or nb.endswith(na):
        return True
    return False


# ---------------------------------------------------------------------------
# Table-type detectors
# ---------------------------------------------------------------------------

def _is_requirement_table(table) -> bool:
    headers = _table_header_row(table)
    return any('requirement' in h and '#' in h for h in headers) or \
           any(h.strip() == 'requirement #' for h in headers)


def _is_specification_table(table) -> bool:
    headers = _table_header_row(table)
    return any('specification' in h and '#' in h for h in headers) or \
           any(h.strip() == 'specification #' for h in headers)


def _is_procedure_table(table) -> bool:
    """
    Identify the Test Procedure table (Section 9) strictly by requiring
    'procedure step description' in a column header.

    This intentionally excludes the Pre-Test Setup table, whose header
    reads 'Setup Step Description' (no 'procedure' keyword).  If no table
    has that exact phrase, fall back to the largest table that has both
    'step id' and 'expected result' – that will be the long procedure table,
    not the short setup table.
    """
    headers = _table_header_row(table)
    # Strict match: proper procedure table header
    if any('procedure step description' in h for h in headers):
        return True
    return False


def _procedure_table_fallback(tables) -> object:
    """
    If strict detection fails, return the table that has Step ID +
    Expected Result AND the most data rows (excludes short setup tables).
    Returns None if no suitable table is found.
    """
    best = None
    best_rows = 0
    for table in tables:
        headers = _table_header_row(table)
        has_step     = any('step id' in h for h in headers)
        has_expected = any('expected result' in h for h in headers)
        if has_step and has_expected and len(table.rows) > best_rows:
            best = table
            best_rows = len(table.rows)
    return best


# ---------------------------------------------------------------------------
# Table parsers
# ---------------------------------------------------------------------------

def _parse_req_spec_table(table) -> List[RequirementRow]:
    """Parse a requirements or specifications table into RequirementRow list."""
    rows = table.rows
    if not rows:
        return []

    headers = _table_header_row(table)

    # Locate column indices
    num_col = _find_col_any(headers, 'requirement #', 'specification #')
    if num_col == -1:
        # Fallback: first column containing SRS/TBR in its header
        num_col = _find_col_any(headers, 'requirement', 'specification')
    if num_col == -1:
        num_col = 0

    desc_col = _find_col(headers, 'description')
    if desc_col == -1 or desc_col == num_col:
        # Pick the non-num, non-ac column
        desc_col = next((i for i in range(len(headers)) if i != num_col), 1)

    ac_col = _find_col(headers, 'acceptance criteria')
    if ac_col == -1:
        ac_col = _find_col(headers, 'acceptance')
    if ac_col == -1:
        # Last column
        ac_col = len(headers) - 1

    result: List[RequirementRow] = []
    for row in rows[1:]:  # skip header row
        cells = row.cells
        if len(cells) <= max(num_col, desc_col, ac_col):
            continue

        number = _cell_text(cells[num_col], " ").strip()

        # Skip rows that don't look like an ID
        if not number or not REQ_NUMBER_RE.search(number):
            continue

        # Use the matched number (in case there's surrounding whitespace/noise)
        matched = REQ_NUMBER_RE.findall(number)
        if matched:
            number = matched[0]

        description = _cell_text(cells[desc_col]) if desc_col < len(cells) else ""
        ac = _cell_text(cells[ac_col]) if ac_col < len(cells) else ""

        result.append(RequirementRow(
            number=number.strip(),
            description=description.strip(),
            existing_ac=ac.strip(),
        ))

    return result


def _parse_procedure_table(table) -> List[ProcedureStep]:
    """Parse the Section 9 test-procedure table into ProcedureStep list."""
    rows = table.rows
    if not rows:
        return []

    headers = _table_header_row(table)

    step_col   = _find_col_any(headers, 'step id', 'step #', 'step')
    desc_col   = _find_col(headers, 'procedure step description')
    if desc_col == -1:
        desc_col = _find_col(headers, 'description')
    result_col = _find_col(headers, 'expected result')
    notes_col  = _find_col(headers, 'notes', 'note')

    # Positional fallbacks
    if step_col   == -1: step_col   = 0
    if desc_col   == -1: desc_col   = 1
    if result_col == -1: result_col = 2
    if notes_col  == -1: notes_col  = 3

    steps: List[ProcedureStep] = []
    for row in rows[1:]:
        cells = row.cells
        if not cells:
            continue

        def get(idx: int) -> str:
            return _cell_text(cells[idx]) if idx < len(cells) else ""

        step_id = get(step_col).strip()
        # Only process rows whose step ID starts with a digit
        if not step_id or not step_id[0].isdigit():
            continue

        desc     = get(desc_col)
        expected = get(result_col)
        notes    = get(notes_col)

        # Extract requirement/spec numbers from Notes, then fall back to
        # Expected Result and Description if Notes is empty or just "N/A"
        ref_numbers = _extract_numbers(notes)
        if not ref_numbers:
            ref_numbers = _extract_numbers(expected) + _extract_numbers(desc)

        # Deduplicate while preserving order
        seen: set = set()
        unique_refs: List[str] = []
        for r in ref_numbers:
            key = normalize(r)
            if key not in seen:
                seen.add(key)
                unique_refs.append(r)

        steps.append(ProcedureStep(
            step_id=step_id,
            description=desc,
            expected_result=expected,
            notes=notes,
            referenced_numbers=unique_refs,
        ))

    return steps


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_document(docx_path: str) -> ParseResult:
    """
    Parse a test-case .docx file and return structured data.

    Searches all tables in document order; identifies each table by its
    column headers rather than by section number, making the parser robust
    to document renumbering.
    """
    result = ParseResult()

    try:
        doc = Document(docx_path)
    except Exception as exc:
        result.errors.append(f"Cannot open document: {exc}")
        return result

    req_found  = False
    spec_found = False
    proc_found = False

    all_tables = doc.tables  # keep reference for fallback search

    for table in all_tables:
        headers = _table_header_row(table)
        if not headers:
            continue

        if not req_found and _is_requirement_table(table):
            rows = _parse_req_spec_table(table)
            if rows:
                result.requirements = rows
                req_found = True
            continue

        if not spec_found and _is_specification_table(table):
            rows = _parse_req_spec_table(table)
            if rows:
                result.specifications = rows
                spec_found = True
            continue

        if not proc_found and _is_procedure_table(table):
            steps = _parse_procedure_table(table)
            if steps:
                result.procedure_steps = steps
                proc_found = True

    # Fallback: find procedure table by row count if strict detection missed it
    if not proc_found:
        fallback = _procedure_table_fallback(all_tables)
        if fallback is not None:
            steps = _parse_procedure_table(fallback)
            if steps:
                result.procedure_steps = steps
                proc_found = True
                result.warnings.append(
                    "Test Procedure table identified via fallback detection "
                    "(no 'Procedure Step Description' column found); "
                    "results may be less accurate."
                )

    # Accumulate warnings for missing sections
    if not req_found:
        result.warnings.append(
            "Requirements table not found. "
            "Expected a table with a 'Requirement #' column."
        )
    if not spec_found:
        result.warnings.append(
            "Specifications table not found. "
            "Expected a table with a 'Specification #' column."
        )
    if not proc_found:
        result.errors.append(
            "Test Procedure table not found. "
            "Expected a table with columns 'Step ID' and 'Expected Result'."
        )

    return result
