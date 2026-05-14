"""
exporter.py – Export acceptance criteria results to Excel or Word.

Excel export:  one worksheet with a row per requirement/spec.
Word export:   a formatted document grouped by type (Requirements / Specifications).
"""

import io
from typing import Dict, Any

from openpyxl import Workbook
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Excel export
# ---------------------------------------------------------------------------

_HEADER_FILL   = PatternFill(fill_type="solid", fgColor="1F4E79")
_REQ_ROW_FILL  = PatternFill(fill_type="solid", fgColor="DEEAF1")
_SPEC_ROW_FILL = PatternFill(fill_type="solid", fgColor="E2EFDA")
_ALT_FILL      = PatternFill(fill_type="solid", fgColor="F5F5F5")
_THIN_BORDER   = Border(
    left=Side(style='thin', color='CCCCCC'),
    right=Side(style='thin', color='CCCCCC'),
    top=Side(style='thin', color='CCCCCC'),
    bottom=Side(style='thin', color='CCCCCC'),
)
_WRAP = Alignment(wrap_text=True, vertical='top')
_CENTRE = Alignment(horizontal='center', vertical='center', wrap_text=True)


def export_to_excel(results: Dict[str, Any]) -> bytes:
    """
    Build an Excel workbook from *results* and return the raw bytes.

    Columns: #  |  Type  |  ID  |  Description  |  Generated AC  |  Verified by Steps
    A second sheet lists every matched step per item for traceability.
    """
    wb = Workbook()

    # ---- Sheet 1: Summary ----
    ws = wb.active
    ws.title = "Acceptance Criteria"

    col_headers = [
        "#",
        "Type",
        "Requirement / Spec #",
        "Description",
        "Generated Acceptance Criteria",
        "Verified by Steps",
    ]

    # Header row
    for col, h in enumerate(col_headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF", size=11, name="Calibri")
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTRE
        cell.border = _THIN_BORDER

    # Freeze header
    ws.freeze_panes = "A2"

    # Data rows
    for idx, (num, item) in enumerate(results.items(), 1):
        row = idx + 1
        row_fill = _REQ_ROW_FILL if item['item_type'] == 'Requirement' else _SPEC_ROW_FILL
        step_ids = ", ".join(f"Step {s['step_id']}" for s in item['matching_steps'])

        values = [
            idx,
            item['item_type'],
            item['number'],
            item['description'],
            item['generated_ac'] if item['generated_ac'] else "(No matching steps found)",
            step_ids or "—",
        ]

        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill = row_fill
            cell.border = _THIN_BORDER
            if col in (1, 2):
                cell.alignment = Alignment(horizontal='center', vertical='top', wrap_text=True)
            else:
                cell.alignment = _WRAP
            cell.font = Font(size=10, name="Calibri")

        # Highlight unmatched rows
        if item['unmatched']:
            for col in range(1, len(col_headers) + 1):
                ws.cell(row=row, column=col).font = Font(
                    size=10, name="Calibri", italic=True, color="999999"
                )

    # Column widths
    widths = [5, 14, 22, 50, 75, 20]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ---- Sheet 2: Step Traceability ----
    ws2 = wb.create_sheet("Step Traceability")
    trace_headers = ["Requirement/Spec #", "Step ID", "Procedure Step Description", "Expected Result"]
    for col, h in enumerate(trace_headers, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF", size=11, name="Calibri")
        cell.fill = _HEADER_FILL
        cell.alignment = _CENTRE
        cell.border = _THIN_BORDER
    ws2.freeze_panes = "A2"

    trace_row = 2
    for num, item in results.items():
        for step in item['matching_steps']:
            ws2.cell(row=trace_row, column=1, value=item['number']).border = _THIN_BORDER
            ws2.cell(row=trace_row, column=2, value=f"Step {step['step_id']}").border = _THIN_BORDER
            ws2.cell(row=trace_row, column=3, value=step['description']).border = _THIN_BORDER
            ws2.cell(row=trace_row, column=4, value=step['expected_result']).border = _THIN_BORDER
            for col in range(1, 5):
                ws2.cell(row=trace_row, column=col).alignment = _WRAP
                ws2.cell(row=trace_row, column=col).font = Font(size=10, name="Calibri")
            trace_row += 1

    for i, w in enumerate([22, 10, 60, 60], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour via XML (python-docx lacks a direct API)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_table_row(table, label: str, value: str, label_color="1F4E79"):
    """Add a two-cell row to a Word table."""
    row = table.add_row()
    # Label cell
    lc = row.cells[0]
    lc.text = label
    _set_cell_bg(lc, "DEEAF1")
    for para in lc.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    # Value cell
    vc = row.cells[1]
    vc.text = value
    for para in vc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
    return row


def export_to_docx(results: Dict[str, Any]) -> bytes:
    """
    Build a Word document summarising acceptance criteria and return raw bytes.
    """
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title = doc.add_heading("Acceptance Criteria Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Stats paragraph
    req_count  = sum(1 for v in results.values() if v['item_type'] == 'Requirement')
    spec_count = sum(1 for v in results.values() if v['item_type'] == 'Specification')
    matched    = sum(1 for v in results.values() if not v['unmatched'])
    stats_para = doc.add_paragraph(
        f"Requirements: {req_count}   |   Specifications: {spec_count}   |   "
        f"Matched to test steps: {matched} / {len(results)}"
    )
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in stats_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()

    # Group by type
    for item_type in ("Requirement", "Specification"):
        type_items = [(k, v) for k, v in results.items() if v['item_type'] == item_type]
        if not type_items:
            continue

        type_heading = doc.add_heading(f"{item_type}s", level=1)

        for num, item in type_items:
            # Sub-heading per item
            h = doc.add_heading(item['number'], level=2)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            # Details table
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'
            tbl.columns[0].width = Cm(4)
            tbl.columns[1].width = Cm(13.5)

            _add_table_row(tbl, "Type", item_type)
            _add_table_row(tbl, "Description", item['description'] or "—")

            # Generated AC row (highlighted)
            ac_row = tbl.add_row()
            lc = ac_row.cells[0]
            lc.text = "Generated AC"
            _set_cell_bg(lc, "1F4E79")
            for para in lc.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)

            vc = ac_row.cells[1]
            ac_text = item['generated_ac'] or "(No matching test steps found)"
            vc.text = ac_text
            _set_cell_bg(vc, "E7F3FF")
            for para in vc.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

            # Steps row
            steps_str = ", ".join(f"Step {s['step_id']}" for s in item['matching_steps']) or "—"
            _add_table_row(tbl, "Test Steps", steps_str)

            doc.add_paragraph()  # spacer between items

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
